#!/usr/bin/env python3
"""
Shift Parser for Monthly Doctor Schedules
Parses DOCX files containing cardiology attending and resident shifts
"""

import os
import json
from datetime import datetime, date
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict
from docx import Document


@dataclass
class DailyShift:
    """Represents shifts for a single day"""
    day: int
    month_name: str
    weekday: str

    # Attending cardiologists
    attendings: List[str]

    # Residents
    major_shift: Optional[str] = None  # ΜΕΓΑΛΕΣ ΕΦΗΜΕΡΙΕΣ 24ωρη
    minor_shift: Optional[str] = None  # ΜΙΚΡΕΣ ΕΦΗΜΕΡΙΕΣ 24ωρη
    tep_cardiologist: Optional[str] = None  # ΚΑΡΔΙΟΛΟΓΟΣ ΤΕΠ 12ωρη

    # Cardiac Surgeons
    senior_cardiac_surgeon: Optional[str] = None  # Ανώτερος Καρδιοχειρουργός
    junior_cardiac_surgeon: Optional[str] = None  # Νεώτερος Καρδιοχειρουργός

    # Anesthesiologists
    anesthesiologist_1: Optional[str] = None  # Αναισθησιολόγος 1
    anesthesiologist_2: Optional[str] = None  # Αναισθησιολόγος 2

    # Pediatric Cardiology
    pediatric_cardiologist: Optional[str] = None  # Παιδοκαρδιολόγος

    def __str__(self):
        """String representation for display"""
        attending_str = ", ".join(self.attendings) if self.attendings else "Κανένας"
        parts = [f"Επιμελητές: {attending_str}"]

        if self.major_shift:
            parts.append(f"Μεγάλη Εφημερία: {self.major_shift}")
        if self.minor_shift:
            parts.append(f"Μικρή Εφημερία: {self.minor_shift}")
        if self.tep_cardiologist:
            parts.append(f"ΤΕΠ: {self.tep_cardiologist}")
        if self.senior_cardiac_surgeon:
            parts.append(f"Καρδιοχειρουργός 1: {self.senior_cardiac_surgeon}")
        if self.junior_cardiac_surgeon:
            parts.append(f"Καρδιοχειρουργός 2: {self.junior_cardiac_surgeon}")
        if self.anesthesiologist_1:
            parts.append(f"Αναισθησιολόγος 1: {self.anesthesiologist_1}")
        if self.anesthesiologist_2:
            parts.append(f"Αναισθησιολόγος 2: {self.anesthesiologist_2}")
        if self.pediatric_cardiologist:
            parts.append(f"Παιδοκαρδιολόγος: {self.pediatric_cardiologist}")

        return " | ".join(parts)


class ShiftParser:
    """Parses monthly shift schedules from DOCX files"""

    # Greek month names mapping
    GREEK_MONTHS = {
        "ΙΑΝΟΥΑΡΙΟΣ": 1, "ΙΑΝΟΥΑΡΙΟΥ": 1,
        "ΦΕΒΡΟΥΑΡΙΟΣ": 2, "ΦΕΒΡΟΥΑΡΙΟΥ": 2,
        "ΜΑΡΤΙΟΣ": 3, "ΜΑΡΤΙΟΥ": 3,
        "ΑΠΡΙΛΙΟΣ": 4, "ΑΠΡΙΛΙΟΥ": 4,
        "ΜΑΪΟΣ": 5, "ΜΑΙΟΥ": 5,
        "ΙΟΥΝΙΟΣ": 6, "ΙΟΥΝΙΟΥ": 6,
        "ΙΟΥΛΙΟΣ": 7, "ΙΟΥΛΙΟΥ": 7,
        "ΑΥΓΟΥΣΤΟΣ": 8, "ΑΥΓΟΥΣΤΟΥ": 8,
        "ΣΕΠΤΕΜΒΡΙΟΣ": 9, "ΣΕΠΤΕΜΒΡΙΟΥ": 9,
        "ΟΚΤΩΒΡΙΟΣ": 10, "ΟΚΤΩΒΡΙΟΥ": 10,
        "ΝΟΕΜΒΡΙΟΣ": 11, "ΝΟΕΜΒΡΙΟΥ": 11,
        "ΔΕΚΕΜΒΡΙΟΣ": 12, "ΔΕΚΕΜΒΡΙΟΥ": 12
    }

    def __init__(self, docx_path: str):
        """Initialize parser with DOCX file path"""
        self.docx_path = docx_path
        self.doc = None
        self.month_number: Optional[int] = None
        self.year: Optional[int] = None
        self.shifts: Dict[int, DailyShift] = {}  # day -> DailyShift

    def parse(self) -> bool:
        """
        Parse the DOCX file and extract all shifts
        Returns True if successful, False otherwise
        """
        try:
            self.doc = Document(self.docx_path)

            # Extract month and year from document
            if not self._extract_month_year():
                print("✗ Αδυναμία εξαγωγής μήνα/έτους από το αρχείο")
                return False

            if len(self.doc.tables) < 2:
                print(f"✗ Αναμενόταν τουλάχιστον 2 πίνακες, βρέθηκαν {len(self.doc.tables)}")
                return False

            # Parse attending cardiologists (Table 1)
            self._parse_attending_table(self.doc.tables[0])

            # Parse residents (Table 2)
            self._parse_resident_table(self.doc.tables[1])

            print(f"✓ Επιτυχής ανάλυση: {len(self.shifts)} ημέρες για {self.month_number}/{self.year}")
            return True

        except Exception as e:
            print(f"✗ Σφάλμα κατά την ανάλυση: {e}")
            return False

    def _extract_month_year(self) -> bool:
        """Extract month and year from document paragraphs"""
        # Look for patterns like "ΟΚΤΩΒΡΙΟΣ 2025"
        for para in self.doc.paragraphs:
            text = para.text.strip().upper()

            # Try to find month name and year
            for month_name, month_num in self.GREEK_MONTHS.items():
                if month_name in text:
                    self.month_number = month_num

                    # Extract year (4 consecutive digits)
                    import re
                    year_match = re.search(r'\b(20\d{2})\b', text)
                    if year_match:
                        self.year = int(year_match.group(1))
                        return True

        return False

    def _parse_attending_table(self, table):
        """Parse attending cardiologists table (first table)"""
        # Table structure:
        # Row 0+: [Day, Month, Weekday, Names]
        # Names can be multi-line, separated by newlines, with * for readiness

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]

            if len(cells) < 4:
                continue

            # Extract day number
            try:
                day = int(cells[0])
            except (ValueError, IndexError):
                continue

            month_name = cells[1]
            weekday = cells[2]
            names_raw = cells[3]

            # Parse names (split by newlines, ignore names with asterisks)
            names = []
            for name in names_raw.split('\n'):
                name = name.strip()
                # Skip names that contain asterisks
                if name and '*' not in name:
                    names.append(name)

            # Create or update shift entry
            if day not in self.shifts:
                self.shifts[day] = DailyShift(
                    day=day,
                    month_name=month_name,
                    weekday=weekday,
                    attendings=names
                )
            else:
                self.shifts[day].attendings = names

    def _parse_resident_table(self, table):
        """Parse residents table (second table)"""
        # Table structure:
        # Row 0: Headers [Day, Month, Weekday, ΜΕΓΑΛΕΣ, ΜΙΚΡΕΣ, ΤΕΠ, ΑΝΑΙΣΘ 1, ΑΝΑΙΣΘ 2, ΠΑΙΔΟΚΑΡΔ]
        # Row 1+: [Day, Month, Weekday, MajorName, MinorName, TEPName, Anest1, Anest2, PedCardio]

        for i, row in enumerate(table.rows):
            if i == 0:  # Skip header row
                continue

            cells = [cell.text.strip() for cell in row.cells]

            if len(cells) < 6:
                continue

            # Extract day number
            try:
                day = int(cells[0])
            except (ValueError, IndexError):
                continue

            month_name = cells[1]
            weekday = cells[2]
            major_shift = cells[3] if cells[3] else None
            minor_shift = cells[4] if cells[4] else None
            tep_cardiologist = cells[5] if cells[5] else None

            # Optional columns for anesthesiologists and pediatric cardiology
            anesthesiologist_1 = cells[6] if len(cells) > 6 and cells[6] else None
            anesthesiologist_2 = cells[7] if len(cells) > 7 and cells[7] else None
            pediatric_cardiologist = cells[8] if len(cells) > 8 and cells[8] else None

            # Update existing shift or create new one
            if day in self.shifts:
                self.shifts[day].major_shift = major_shift
                self.shifts[day].minor_shift = minor_shift
                self.shifts[day].tep_cardiologist = tep_cardiologist
                self.shifts[day].anesthesiologist_1 = anesthesiologist_1
                self.shifts[day].anesthesiologist_2 = anesthesiologist_2
                self.shifts[day].pediatric_cardiologist = pediatric_cardiologist
            else:
                self.shifts[day] = DailyShift(
                    day=day,
                    month_name=month_name,
                    weekday=weekday,
                    attendings=[],
                    major_shift=major_shift,
                    minor_shift=minor_shift,
                    tep_cardiologist=tep_cardiologist,
                    anesthesiologist_1=anesthesiologist_1,
                    anesthesiologist_2=anesthesiologist_2,
                    pediatric_cardiologist=pediatric_cardiologist
                )

    def get_shift_for_day(self, day: int) -> Optional[DailyShift]:
        """Get shift information for a specific day"""
        return self.shifts.get(day)

    def get_shift_for_date(self, target_date: date) -> Optional[DailyShift]:
        """Get shift information for a specific date"""
        if target_date.month != self.month_number or target_date.year != self.year:
            return None

        return self.shifts.get(target_date.day)

    def validate_month_year(self, expected_month: int, expected_year: int) -> bool:
        """Check if parsed file matches expected month and year"""
        return self.month_number == expected_month and self.year == expected_year

    def save_to_json(self, filepath: str = "shifts_cache.json"):
        """Save parsed shifts to JSON for caching"""
        data = {
            'month': self.month_number,
            'year': self.year,
            'last_update': datetime.now().isoformat(),
            'shifts': {
                day: asdict(shift) for day, shift in self.shifts.items()
            }
        }

        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        print(f"✓ Εφημερίες αποθηκεύτηκαν στο {filepath}")

    @classmethod
    def load_from_json(cls, filepath: str = "shifts_cache.json") -> Optional['ShiftParser']:
        """Load shifts from cached JSON file"""
        if not os.path.exists(filepath):
            return None

        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Create a parser instance without file
            parser = cls.__new__(cls)
            parser.docx_path = None
            parser.doc = None
            parser.month_number = data['month']
            parser.year = data['year']
            parser.shifts = {}

            # Reconstruct DailyShift objects
            for day_str, shift_data in data['shifts'].items():
                day = int(day_str)
                parser.shifts[day] = DailyShift(**shift_data)

            print(f"✓ Εφημερίες φορτώθηκαν από {filepath}")
            return parser

        except Exception as e:
            print(f"✗ Σφάλμα φόρτωσης cache: {e}")
            return None

    def update_shift(self, day: int, field: str, value: str):
        """
        Update a specific field in a shift
        field can be: 'attendings', 'major_shift', 'minor_shift', 'tep_cardiologist',
                      'senior_cardiac_surgeon', 'junior_cardiac_surgeon',
                      'anesthesiologist_1', 'anesthesiologist_2', 'pediatric_cardiologist'
        """
        if day not in self.shifts:
            print(f"✗ Δεν βρέθηκε εφημερία για την ημέρα {day}")
            return False

        shift = self.shifts[day]

        if field == 'attendings':
            # Parse comma-separated names
            shift.attendings = [name.strip() for name in value.split(',') if name.strip()]
        elif field in ['major_shift', 'minor_shift', 'tep_cardiologist',
                       'senior_cardiac_surgeon', 'junior_cardiac_surgeon',
                       'anesthesiologist_1', 'anesthesiologist_2', 'pediatric_cardiologist']:
            setattr(shift, field, value.strip() if value.strip() else None)
        else:
            print(f"✗ Άγνωστο πεδίο: {field}")
            return False

        print(f"✓ Ενημερώθηκε η εφημερία της {day}/{self.month_number}")
        return True


def test_parser():
    """Test function to demonstrate parser usage"""
    import sys

    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
    else:
        docx_path = "/home/dimitris/Έγγραφα/OCSC TEP display/ΕΦΗΜΕΡΙΕΣ ΟΚΤΩΒΡΙΟΣ 2025.docx"

    if not os.path.exists(docx_path):
        print(f"✗ Το αρχείο δεν βρέθηκε: {docx_path}")
        return

    print(f"Ανάλυση: {docx_path}\n")

    parser = ShiftParser(docx_path)

    if parser.parse():
        print(f"\n{'='*60}")
        print(f"Εφημερίες {parser.month_number}/{parser.year}")
        print(f"{'='*60}\n")

        # Show today's shift if available
        today = date.today()
        today_shift = parser.get_shift_for_date(today)

        if today_shift:
            print(f"Σημερινή εφημερία ({today.day}/{today.month}):")
            print(f"  {today_shift}\n")

        # Show first 5 days
        print("Πρώτες 5 ημέρες:")
        for day in sorted(parser.shifts.keys())[:5]:
            shift = parser.shifts[day]
            print(f"\n{shift.day} {shift.month_name} ({shift.weekday}):")
            print(f"  {shift}")

        # Test saving and loading
        parser.save_to_json()


if __name__ == "__main__":
    test_parser()
